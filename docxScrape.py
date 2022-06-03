from docx import Document
from lxml import etree
import zipfile
from pathlib import Path
import json
from bs4 import BeautifulSoup
import re


#taken from https://stackoverflow.com/questions/47390928/extract-docx-comments
ooXMLns = {'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
#Function to extract all the comments of document(Same as accepted answer)
#Returns a dictionary with comment id as key and comment string as value
def get_document_comments(docxFileName):
  comments_dict={}
  docxZip = zipfile.ZipFile(docxFileName)
  commentsXML = docxZip.read('word/comments.xml')
  et = etree.XML(commentsXML)
  comments = et.xpath('//w:comment',namespaces=ooXMLns)
  # documentXML = docxZip.read('word/document.xml')





  # root = etree.parse(documentXML).getroot()
  # instruments = root.find('commentRangeStart')
  # instrument = instruments.findall('commentRangeStart')
  # for grandchild in instrument:
  #     code, source = grandchild.find('Code'), grandchild.find('Source')
  #     print (code.text), (source.text)


  #doc_et = etree.XML(documentXML)
  #print(etree.tostring(doc_et, pretty_print=True))
  linked_texts = et.xpath('@w:id=\"112\"', namespaces=ooXMLns)
  for c in comments:
    comment=c.xpath('string(.)',namespaces=ooXMLns)
    comment_id=c.xpath('@w:id',namespaces=ooXMLns)[0]
    comments_dict[comment_id]=comment
  return comments_dict

#Function to fetch all the comments in a paragraph
def paragraph_comments(paragraph,comments_dict):
  comments=[]
  for run in paragraph.runs:
    comment_reference=run._r.xpath("./w:commentReference")
    if comment_reference:
      comment_id=comment_reference[0].xpath('@w:id',namespaces=ooXMLns)[0]
      comment=comments_dict[comment_id]
      comments.append(comment)
  return comments

#Function to fetch all comments with their referenced paragraph
#This will return list like this [{'Paragraph text': [comment 1,comment 2]}]
def comments_with_reference_paragraph(docxFileName):
  document = Document(docxFileName)
  comments_dict=get_document_comments(docxFileName)
  comments_with_their_reference_paragraph=[]




  # for paragraph in document.paragraphs:
    #print(paragraph._p.xml)

  docxZip = zipfile.ZipFile(docxFileName)
  documentXML = docxZip.read('word/document.xml')

  with docxZip.open('word/document.xml') as fp:
    soup = BeautifulSoup(fp.read(), 'lxml')
    print(soup.contents)
    linked_comments = re.findall(r'<w:commentrangestart[\S\s]+?(?=<w:commentrangeend)', str(soup.contents))

    

  # print(soup.contents)


    
  # soup = BeautifulSoup(document.paragraphs._p.xml, 'lxml')

  # print(soup.contents)
  # print(type(soup.contents))


  # print(soup.find('w:p'))
  # print(soup.find('w:commentRangeStart'))






  #   if comments_dict: 
  #     comments=paragraph_comments(paragraph,comments_dict)  
  #     if comments:
  #         comments_with_their_reference_paragraph.append({paragraph.text: comments})
  # return comments_with_their_reference_paragraph

  
'''
observations:

- use paragraph._p.xml to scrape xml
- located inside <w:r><w:t>TEXT</w:t></w:r>
- in between <w:commentRangeStart w:id="{num}"> <w:commentRangeEnd w:id="{num}">

'''


if __name__=="__main__":
  articles_path = Path('~/Documents/Research/MediaArticles/')
  document = articles_path / 'AlabamaArticles.docx'  #filepath for the input document
  comments_with_reference_paragraph('AlabamaArticles.docx')